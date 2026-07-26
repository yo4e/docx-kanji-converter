from pathlib import Path

import pytest
from docx import Document

from main import default_output_path, main


def test_default_output_path():
    assert default_output_path(Path("novel.docx")) == Path("novel.converted.docx")


def test_cli_dry_run(tmp_path, capsys):
    source = tmp_path / "input.docx"
    Document().save(source)

    assert main([str(source), "--dry-run"]) == 0
    output = capsys.readouterr().out
    assert "Dry run" in output
    assert not (tmp_path / "input.converted.docx").exists()


def test_cli_refuses_to_overwrite_input(tmp_path):
    source = tmp_path / "input.docx"
    Document().save(source)

    with pytest.raises(SystemExit) as error:
        main([str(source), str(source)])

    assert error.value.code == 2


def test_cli_refuses_existing_output_without_force(tmp_path):
    source = tmp_path / "input.docx"
    output = tmp_path / "output.docx"
    Document().save(source)
    Document().save(output)

    with pytest.raises(SystemExit) as error:
        main([str(source), str(output)])

    assert error.value.code == 2


def test_cli_can_replace_existing_output_with_force(tmp_path):
    source = tmp_path / "input.docx"
    output = tmp_path / "output.docx"
    source_doc = Document()
    source_doc.add_paragraph("本文")
    source_doc.save(source)
    Document().save(output)

    assert main([str(source), str(output), "--force"]) == 0
    assert Document(output).paragraphs[0].text == "　本文"


def test_cli_reports_invalid_docx_without_traceback(tmp_path, capsys):
    source = tmp_path / "broken.docx"
    source.write_text("not a docx", encoding="utf-8")

    with pytest.raises(SystemExit) as error:
        main([str(source), "--dry-run"])

    assert error.value.code == 1
    assert "error:" in capsys.readouterr().err
