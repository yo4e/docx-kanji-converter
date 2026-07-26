from pathlib import Path

from docx import Document
from docx.shared import Pt

from converter import ConversionOptions, convert_document, process_document


def test_process_document_applies_rules_and_preserves_run_boundaries():
    document = Document()
    paragraph = document.add_paragraph()
    first = paragraph.add_run("ABC 22...")
    first.italic = True
    paragraph.add_run("本当！")

    heading = document.add_heading("Heading 12", level=1)

    report = process_document(document)

    assert paragraph.text == "　ＡＢＣ 二十二……本当！　"
    assert len(paragraph.runs) == 2
    assert first.italic is False
    assert first.bold is True
    assert first.font.size == Pt(12)
    assert heading.text == "Ｈｅａｄｉｎｇ 十二"
    assert not heading.text.startswith("　")
    assert report.total_changes > 0


def test_process_document_is_idempotent_for_indentation_and_punctuation():
    document = Document()
    paragraph = document.add_paragraph("本文！")

    process_document(document)
    first_result = paragraph.text
    process_document(document)

    assert paragraph.text == first_result


def test_disabled_rule_is_not_applied():
    document = Document()
    paragraph = document.add_paragraph("ABC 22")
    options = ConversionOptions.with_disabled(["ascii", "numbers", "font"])

    process_document(document, options)

    assert paragraph.text == "　ABC 22"


def test_large_number_is_reported_and_not_changed():
    document = Document()
    paragraph = document.add_paragraph("12345円")

    report = process_document(document)

    assert paragraph.text == "　12345円"
    assert report.skipped["numbers_over_4_digits"] == 1


def test_table_and_header_are_left_unchanged():
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "ABC 22"
    document.sections[0].header.paragraphs[0].text = "Header 12"

    process_document(document)

    assert table.cell(0, 0).text == "ABC 22"
    assert document.sections[0].header.paragraphs[0].text == "Header 12"


def test_text_split_across_runs_is_not_joined_for_replacement():
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run(".")
    paragraph.add_run("..")

    process_document(document)

    assert paragraph.text == "　..."


def test_large_bold_normal_paragraph_is_treated_as_heading():
    document = Document()
    heading = document.add_paragraph()
    heading.add_run("第")
    heading.add_run("1")
    heading.add_run("章：四つの目覚め")
    for run in heading.runs:
        run.bold = True
        run.font.name = "ArialUnicodeMS"
        run.font.size = Pt(17)

    process_document(document)

    assert heading.text == "第一章：四つの目覚め"
    assert not heading.text.startswith("　")
    for run in heading.runs:
        assert run.font.name == "ArialUnicodeMS"
        assert run.font.size == Pt(17)


def test_bold_13pt_section_number_is_treated_as_heading():
    document = Document()
    heading = document.add_paragraph()
    run = heading.add_run("2")
    run.bold = True
    run.font.name = "Arial-BoldMT"
    run.font.size = Pt(13)

    process_document(document)

    assert heading.text == "二"
    assert run.font.name == "Arial-BoldMT"
    assert run.font.size == Pt(13)


def test_bold_body_text_below_13pt_is_not_inferred_as_heading():
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("重要な本文")
    run.bold = True
    run.font.name = "ArialUnicodeMS"
    run.font.size = Pt(11)

    process_document(document)

    assert paragraph.text == "　重要な本文"
    assert run.font.name is None
    assert run.font.size == Pt(12)


def test_percent_is_spelled_out_before_number_conversion():
    document = Document()
    paragraph = document.add_paragraph("出力30%に上昇")

    report = process_document(document)

    assert paragraph.text == "　出力三十パーセントに上昇"
    assert report.changes["replacements"] == 1
    assert report.changes["numbers"] == 1


def test_replacements_rule_can_be_disabled():
    document = Document()
    paragraph = document.add_paragraph("出力30%")
    options = ConversionOptions.with_disabled(["replacements"])

    process_document(document, options)

    assert paragraph.text == "　出力三十%"


def test_convert_document_dry_run_does_not_write_output(tmp_path: Path):
    source = tmp_path / "input.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("本文")
    document.save(source)

    report = convert_document(source, None, dry_run=True)

    assert report.total_changes > 0
    assert not output.exists()


def test_convert_document_writes_separate_output_and_keeps_source(tmp_path: Path):
    source = tmp_path / "input.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("本文")
    document.save(source)

    convert_document(source, output)

    assert Document(source).paragraphs[0].text == "本文"
    assert Document(output).paragraphs[0].text == "　本文"
